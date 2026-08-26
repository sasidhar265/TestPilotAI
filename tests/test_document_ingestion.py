from io import BytesIO
from zipfile import ZipFile

import pytest
from docx import Document
from PIL import Image

from app.services.document_ingestion import DocumentIngestionError, DocumentIngestionService


def test_extracts_docx_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("As a customer, I can reset my password.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Given"
    table.cell(0, 1).text = "a registered email"
    content = BytesIO()
    document.save(content)

    result = DocumentIngestionService().extract("requirements.DOCX", content.getvalue())

    assert result.media_type.endswith("wordprocessingml.document")
    assert "reset my password" in result.text
    assert "Given | a registered email" in result.text


def test_extracts_image_text_through_injected_local_ocr() -> None:
    content = BytesIO()
    Image.new("RGB", (20, 20), "white").save(content, format="PNG")

    result = DocumentIngestionService(ocr=lambda image: "  User can log in  \n\n safely ").extract(
        "story.png", content.getvalue()
    )

    assert result.text == "User can log in\nsafely"


@pytest.mark.parametrize("filename", ["story.doc", "story.xls", "story.txt", "story.exe"])
def test_rejects_unsupported_formats(filename: str) -> None:
    with pytest.raises(DocumentIngestionError, match="Unsupported file type"):
        DocumentIngestionService().extract(filename, b"content")


def test_rejects_empty_and_oversized_uploads() -> None:
    service = DocumentIngestionService(max_file_bytes=3)
    with pytest.raises(DocumentIngestionError, match="empty"):
        service.extract("story.pdf", b"")
    with pytest.raises(DocumentIngestionError, match="exceeds"):
        service.extract("story.pdf", b"four")


def test_rejects_ocr_output_over_generation_limit() -> None:
    content = BytesIO()
    Image.new("RGB", (5, 5), "white").save(content, format="JPEG")

    with pytest.raises(DocumentIngestionError, match="Extracted text exceeds"):
        DocumentIngestionService(max_text_chars=5, ocr=lambda image: "too long").extract(
            "story.jpg", content.getvalue()
        )


@pytest.mark.parametrize(
    ("filename", "media_type"),
    [
        ("requirements.pages", "application/vnd.apple.pages"),
        ("requirements.numbers", "application/vnd.apple.numbers"),
    ],
)
def test_extracts_apple_iwork_quicklook_html(filename: str, media_type: str) -> None:
    content = BytesIO()
    with ZipFile(content, "w") as package:
        package.writestr(
            "QuickLook/Preview.html",
            "<html><style>hidden</style><body>AC-1: Customer can check out</body></html>",
        )

    result = DocumentIngestionService().extract(filename, content.getvalue())

    assert result.media_type == media_type
    assert result.text == "AC-1: Customer can check out"


def test_iwork_binary_only_package_has_export_guidance() -> None:
    content = BytesIO()
    with ZipFile(content, "w") as package:
        package.writestr("Index/Document.iwa", b"binary")

    with pytest.raises(DocumentIngestionError, match="export it as PDF.*Excel/XLSX"):
        DocumentIngestionService().extract("requirements.pages", content.getvalue())
