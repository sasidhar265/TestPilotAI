"""Safe, transport-independent extraction of requirement text from documents."""

from __future__ import annotations

from collections.abc import Callable
from dataclasses import dataclass
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from app.agents import AgentKind, FunctionalAgentDescriptor
from app.models import GenerateRequest, TestFormat


class DocumentIngestionError(ValueError):
    """Raised when an uploaded document cannot be safely converted to text."""


@dataclass(frozen=True)
class ExtractedDocument:
    filename: str
    media_type: str
    text: str


class _VisibleHTMLText(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._hidden_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.casefold() in {"script", "style"}:
            self._hidden_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag.casefold() in {"script", "style"} and self._hidden_depth:
            self._hidden_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self._hidden_depth:
            self.parts.append(data)


class DocumentIngestionService:
    """Extract text while enforcing bounded input and output sizes.

    Image extraction uses local Tesseract OCR. The executable must be installed on
    the host; no document content is sent to an external service.
    """

    SUPPORTED_EXTENSIONS = frozenset(
        {".docx", ".pdf", ".png", ".jpg", ".jpeg", ".xlsx", ".pages", ".numbers"}
    )
    MEDIA_TYPES = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pdf": "application/pdf",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".pages": "application/vnd.apple.pages",
        ".numbers": "application/vnd.apple.numbers",
    }

    def __init__(
        self,
        *,
        max_file_bytes: int = 15 * 1024 * 1024,
        max_text_chars: int = 30_000,
        max_archive_entries: int = 2000,
        max_archive_uncompressed_bytes: int = 60 * 1024 * 1024,
        ocr: Callable[[object], str] | None = None,
    ) -> None:
        self.max_file_bytes = max_file_bytes
        self.max_text_chars = max_text_chars
        self.max_archive_entries = max_archive_entries
        self.max_archive_uncompressed_bytes = max_archive_uncompressed_bytes
        self._ocr = ocr

    def extract(self, filename: str, content: bytes) -> ExtractedDocument:
        suffix = Path(filename).suffix.casefold()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            supported = ", ".join(sorted(self.SUPPORTED_EXTENSIONS))
            raise DocumentIngestionError(f"Unsupported file type. Supported types: {supported}")
        if not content:
            raise DocumentIngestionError("The uploaded document is empty.")
        if len(content) > self.max_file_bytes:
            raise DocumentIngestionError(
                f"The uploaded document exceeds the {self.max_file_bytes}-byte limit."
            )
        if suffix in {".docx", ".xlsx", ".pages", ".numbers"}:
            self._validate_archive(content)

        extractor = {
            ".docx": self._extract_docx,
            ".pdf": self._extract_pdf,
            ".png": self._extract_image,
            ".jpg": self._extract_image,
            ".jpeg": self._extract_image,
            ".xlsx": self._extract_xlsx,
            ".pages": self._extract_iwork,
            ".numbers": self._extract_iwork,
        }[suffix]
        try:
            text = self._normalize(extractor(content))
        except DocumentIngestionError:
            raise
        except Exception as error:
            raise DocumentIngestionError(
                f"Could not read {filename}; the file may be corrupt or password-protected."
            ) from error
        if not text:
            hint = " Run OCR on scanned pages first." if suffix == ".pdf" else ""
            raise DocumentIngestionError(f"No readable text was found in {filename}.{hint}")
        if len(text) > self.max_text_chars:
            raise DocumentIngestionError(
                f"Extracted text exceeds the {self.max_text_chars}-character limit."
            )
        return ExtractedDocument(filename, self.MEDIA_TYPES[suffix], text)

    def _validate_archive(self, content: bytes) -> None:
        """Reject malformed or unexpectedly expansive Office/iWork ZIP containers."""
        try:
            with ZipFile(BytesIO(content)) as archive:
                entries = archive.infolist()
                if len(entries) > self.max_archive_entries:
                    raise DocumentIngestionError("The document archive contains too many entries.")
                expanded_size = sum(entry.file_size for entry in entries)
                if expanded_size > self.max_archive_uncompressed_bytes:
                    raise DocumentIngestionError(
                        "The document expands beyond the safe processing limit."
                    )
                for entry in entries:
                    if entry.file_size > 10 * 1024 * 1024 and entry.compress_size == 0:
                        raise DocumentIngestionError(
                            "The document contains an invalid compressed entry."
                        )
                    if entry.compress_size and entry.file_size / entry.compress_size > 200:
                        raise DocumentIngestionError(
                            "The document contains an unsafe compression ratio."
                        )
        except BadZipFile as error:
            raise DocumentIngestionError("The document is not a valid archive.") from error

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        from docx import Document

        document = Document(BytesIO(content))
        chunks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            chunks.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(chunks)

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as error:
            raise DocumentIngestionError(
                "PDF support is not installed on the server. Install the application dependencies "
                "with pip install -e . and restart it."
            ) from error

        try:
            reader = PdfReader(BytesIO(content), strict=False)
        except Exception as error:
            raise DocumentIngestionError(
                "The PDF structure could not be read. Re-save or print the file "
                "as a new PDF and retry."
            ) from error
        if reader.is_encrypted:
            try:
                unlocked = reader.decrypt("") != 0
            except Exception:
                unlocked = False
            if not unlocked:
                raise DocumentIngestionError(
                    "This PDF is password-protected. Remove the password on your Mac, "
                    "then upload it again."
                )
        try:
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as error:
            raise DocumentIngestionError(
                "The PDF opened, but its text layer could not be extracted. "
                "On your Mac, open it in Preview and choose File > Export as PDF, then retry."
            ) from error

    def _extract_image(self, content: bytes) -> str:
        from PIL import Image

        with Image.open(BytesIO(content)) as image:
            image.verify()
        with Image.open(BytesIO(content)) as image:
            if image.width * image.height > 40_000_000:
                raise DocumentIngestionError("Image dimensions are too large to process safely.")
            if self._ocr is not None:
                return self._ocr(image)
            try:
                import pytesseract

                return str(pytesseract.image_to_string(image))
            except ImportError as error:
                raise DocumentIngestionError("Image OCR support is not installed.") from error
            except Exception as error:
                raise DocumentIngestionError(
                    "Image OCR failed. Ensure the Tesseract executable is installed."
                ) from error

    @staticmethod
    def _extract_xlsx(content: bytes) -> str:
        from openpyxl import load_workbook

        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        chunks: list[str] = []
        try:
            for sheet in workbook.worksheets:
                chunks.append(f"Sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    values = [str(value).strip() for value in row if value is not None]
                    if values:
                        chunks.append(" | ".join(values))
        finally:
            workbook.close()
        return "\n".join(chunks)

    @classmethod
    def _extract_iwork(cls, content: bytes) -> str:
        """Read text exposed by an iWork package without launching desktop apps.

        Current Pages and Numbers documents store their source in proprietary IWA
        streams. Many packages also include a Quick Look PDF/HTML preview; older
        packages expose XML. Those representations are safe to extract locally.
        """
        try:
            with ZipFile(BytesIO(content)) as package:
                names = {name.casefold(): name for name in package.namelist()}
                for candidate in ("quicklook/preview.pdf", "preview.pdf"):
                    if candidate in names:
                        return cls._extract_pdf(package.read(names[candidate]))
                for candidate in ("quicklook/preview.html", "preview.html"):
                    if candidate in names:
                        parser = _VisibleHTMLText()
                        parser.feed(package.read(names[candidate]).decode("utf-8", "replace"))
                        return "\n".join(parser.parts)
                for candidate in ("index.xml", "index.apxl"):
                    if candidate in names:
                        import xml.etree.ElementTree as ET

                        root = ET.fromstring(package.read(names[candidate]))
                        return "\n".join(text for text in root.itertext())
        except BadZipFile as error:
            raise DocumentIngestionError(
                "The Apple iWork file is not a readable Pages or Numbers package."
            ) from error
        raise DocumentIngestionError(
            "This Pages or Numbers file contains only Apple's binary IWA data. "
            "On your Mac, export it as PDF (Pages) or Excel/XLSX (Numbers), then upload the export."
        )

    @staticmethod
    def _normalize(text: str) -> str:
        lines = (" ".join(line.split()) for line in text.replace("\x00", "").splitlines())
        return "\n".join(line for line in lines if line).strip()


class InputAgent:
    """Normalize text and convert extracted documents into generation requests."""

    descriptor = FunctionalAgentDescriptor(
        id="input-agent",
        name="Input Agent",
        kind=AgentKind.INPUT,
        purpose="Normalize pasted requirements and extract requirements from supported documents.",
        runtime="local-deterministic",
        capabilities=("text", "docx", "pdf", "xlsx", "pages", "numbers", "png-ocr", "jpeg-ocr"),
    )

    def __init__(self, documents: DocumentIngestionService | None = None) -> None:
        self.documents = documents or DocumentIngestionService()

    def from_text(
        self,
        description: str,
        additional_context: str = "",
        output_format: TestFormat = TestFormat.NORMAL,
    ) -> GenerateRequest:
        return GenerateRequest(
            description=description,
            additional_context=additional_context,
            output_format=output_format,
        )

    def from_document(
        self,
        filename: str,
        content: bytes,
        additional_context: str = "",
        output_format: TestFormat = TestFormat.NORMAL,
    ) -> tuple[ExtractedDocument, GenerateRequest]:
        document = self.documents.extract(filename, content)
        return document, self.from_text(document.text, additional_context, output_format)
